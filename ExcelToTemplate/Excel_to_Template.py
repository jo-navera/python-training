# Feature 1: change saved filename to include rigname "rigname.docx"
# Feature 2: change process_paragraphs to replace placeholder text with the corresponding value from the datarow
# Feature 3: change process_tables to replace placeholder text with the corresponding value from the datarow
# Feature 4: change parse_arguments to accept command line arguments for source file and template file

import pandas as pd
import argparse
import docx
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from datetime import datetime
import re


def parse_arguments():
    return "ExcelToTemplate\Source.xlsx", "ExcelToTemplate\Template - Copy.docx"


def checkedElement():
    elm = OxmlElement('w:checked')
    elm.set(qn('w:val'), "true")
    return elm


def get_spreadsheet_data(df):
    # Get the data from the DataFrame as a list of lists
    data = df.values.tolist()
    return data


def load_spreadsheet(file_path):
    # Load the spreadsheet into a pandas DataFrame
    df = pd.read_excel(file_path)
    return df


def load_template(file_path):
    # Load the template into a python-docx Document
    doc = Document(file_path)
    return doc


def arrange_data_for_processing(df):
    # Arrange the data from the DataFrame into a list of dictionaries
    data = df.to_dict(orient='records')
    data = [{k.upper(): v for k, v in d.items()} for d in data]
    return data


def process_data_row(datarow, template_copy):
    # Process a single row of data and fill the template with the data
    template_copy = process_paragraphs(datarow, template_copy)
    template_copy = process_tables(datarow, template_copy)
    return template_copy


def process_paragraphs(datarow, template_copy):
    # Process the paragraphs in the template and fill them with the data from the datarow
    for paragraph in template_copy.paragraphs:
        keys = extract_keys(paragraph.text)
        if len(keys) > 0:
            for key in keys:
                if f"<<{key.upper}>>" in paragraph.text.upper():
                    paragraph.text = paragraph.text.replace(
                        f"<<{key.upper}>>",  str(
                            datarow[key]))

    return template_copy


def process_tables(datarow, template_copy):
    # Process the tables in the template and fill them with the data from the datarow
    for table in template_copy.tables:
        for row in table.rows:
            for cell in row.cells:
                keys = extract_keys(cell.text)
                if len(keys) > 0:
                    for key in keys:
                        if key in datarow and f"<<{key.upper()}>>" in cell.text.upper():
                            cell.text = cell.text.replace(f"<<{key}>>", str(
                                datarow[key]))
                else:
                    handle_checkboxes_in_cell(datarow, cell)

    return template_copy


def extract_keys(text):
    pattern = r"<<(.*?)>>"
    regex = re.compile(pattern)

    return re.findall(pattern, text)


def save_filled_template(template_copy, output_file):
    # Save the filled template to a new file
    template_copy.save(output_file)


def process_data(data, template):
    # Process all rows of data and fill the template with the data
    for i, datarow in enumerate(data):
        template = load_template(template_file)
        filled_template = process_data_row(datarow, template)
        now = datetime.now()
        output_file = f"filled_template_{i}_{now.strftime("%Y-%m-%d %H%M%S")}.docx"
        save_filled_template(filled_template, output_file)
        print(f"Saved filled template to {output_file}")


def handle_checkboxes_in_cell(datarow, cell):
    legacy_boxes = cell._element.xpath('.//w:checkBox')
    if len(legacy_boxes) > 0:
        for key, value in datarow.items():
            if check_key_in_checkbox(key, cell.text):
                for box in legacy_boxes:
                    checked_el = box.find(qn('w:checked'))
                    index_to_check = get_index_of_checkbox(
                        key, value, cell.text)
                    if (index_to_check < 0):
                        break
                    if legacy_boxes.index(box) == index_to_check:
                        if checked_el is not None:
                            # '1' or 'true' means checked
                            checked_el.set(qn('w:val'), '1')
                            break
                        else:
                            # Create and append a new <w:checked> tag if missing
                            new_checked = docx.oxml.shared.OxmlElement(
                                'w:checked')
                            new_checked.set(qn('w:val'), '1')
                            box.append(new_checked)
                            break


def check_key_in_checkbox(key, cell_text):
    temp_key = key.replace(" ", "").upper()
    temp_cell_text = cell_text.replace(" ", "").upper()
    if temp_key.startswith("NAMEOF"):
        # If key start with "NameOF", it's value is a string, rather than a list of options
        return False
    return temp_key in temp_cell_text


def check_key_with_checkbox(key, cell_text):
    temp_string = cell_text.replace("\t", "").replace(" ", "").upper()
    val = key.replace(" ", "").upper() in temp_string.upper()
    return val


def get_index_of_checkbox(key, value, cell_text):
    temp_string = cell_text
    options = temp_string.upper().split("\t")
    options[:] = [item.strip() for item in options]
    if not (str(value).upper() in options):
        # Key must be combined with options
        # Use the key to derive options instead
        options = key.split(" OR ")
        if value.upper() not in options:
            return -1
        return options.index(value.upper())
    return (options.index(value.upper()) - 1)


if __name__ == "__main__":
    source_file, template_file = parse_arguments()
    loaded_df = load_spreadsheet(source_file)
    arranged_datarows = arrange_data_for_processing(loaded_df)

    process_data(arranged_datarows, template_file)
